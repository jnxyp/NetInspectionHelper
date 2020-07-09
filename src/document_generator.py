import os
from os.path import join, exists

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm

from config import REPORT_PATH, REPORT_FILENAME, SCREENSHOT_FILENAME_PAGE, \
    SCREENSHOT_FILENAME_SCREEN, REPORT_IMAGE_WIDTH, REPORT_FONT_NAME, SCREENSHOT_TAKE_SCREEN, \
    SCREENSHOT_TAKE_PAGE
from sites import Site, SITES


class InspectionReport:
    def __init__(self, company_name: str, sites: list):
        self.company_name = company_name
        self.sites = sites
        self.document = Document()

    def generate(self):
        d = self.document
        # Set page size
        section = d.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        # Set font
        style = d.styles['List Number']
        style.font.name = REPORT_FONT_NAME
        style._element.rPr.rFonts.set(qn('w:eastAsia'), REPORT_FONT_NAME)
        # Generate contents
        self.generate_header()
        for site in self.sites:
            self.generate_site_report(site)

    def generate_header(self):
        d = self.document
        d.add_heading(self.company_name, 0)

    def generate_site_report(self, site: Site):
        d = self.document
        p = d.add_paragraph(
            site.get_name(), style='List Number'
        )

        screen_shot_path_page = site.get_screenshot_file_name_full(self.company_name,
                                                                   SCREENSHOT_FILENAME_PAGE)
        screen_shot_path_screen = site.get_screenshot_file_name_full(self.company_name,
                                                                     SCREENSHOT_FILENAME_SCREEN)
        if SCREENSHOT_TAKE_PAGE and exists(screen_shot_path_page):
            d.add_picture(screen_shot_path_page, REPORT_IMAGE_WIDTH)
        if SCREENSHOT_TAKE_SCREEN and exists(screen_shot_path_screen):
            d.add_picture(screen_shot_path_screen, REPORT_IMAGE_WIDTH)

        d.add_page_break()

    def save(self, path: str = None):
        if path is None:
            path = join(REPORT_PATH, REPORT_FILENAME.format(company_name=self.company_name))
        if not exists(REPORT_PATH):
            os.makedirs(REPORT_PATH)
        self.document.save(path)

if __name__ == '__main__':
    report = InspectionReport("常高新集团有限公司", SITES)
    report.generate()
    report.save()